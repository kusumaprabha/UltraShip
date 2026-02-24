import os
import uuid
import hashlib
from typing import List, Dict, Any, Optional
import PyPDF2
from docx import Document
import numpy as np
from sentence_transformers import SentenceTransformer
import faiss
import pickle
import json

class DocumentProcessor:
    def __init__(self, upload_dir="uploads", index_dir="vector_indices"):
        self.upload_dir = upload_dir
        self.index_dir = index_dir
        os.makedirs(upload_dir, exist_ok=True)
        os.makedirs(index_dir, exist_ok=True)
        
        # Initialize embedding model (lightweight, runs locally)
        print("🔄 Loading embedding model...")
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        print("✅ Embedding model loaded")
        
        # Store document metadata
        self.documents = {}
        
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"Error extracting PDF: {e}")
        return text
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
        return text
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                print(f"Error extracting TXT: {e}")
                return ""
        except Exception as e:
            print(f"Error extracting TXT: {e}")
            return ""
    
    def extract_text(self, file_path: str, file_extension: str) -> str:
        """Extract text based on file extension"""
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
    
    def clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        return '\n'.join(lines)
    
    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 100) -> List[str]:
        """
        Intelligently chunk text with overlap
        Tries to break at sentence boundaries when possible
        """
        if not text:
            return []
        
        chunks = []
        
        # Split into sentences (simple approach)
        sentences = text.replace('\n', ' ').split('. ')
        
        current_chunk = ""
        current_size = 0
        
        for sentence in sentences:
            sentence = sentence.strip() + ". "
            sentence_size = len(sentence.split())
            
            if current_size + sentence_size <= chunk_size:
                current_chunk += sentence
                current_size += sentence_size
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap from previous chunk
                words = current_chunk.split()
                overlap_words = words[-overlap:] if len(words) > overlap else words
                current_chunk = ' '.join(overlap_words) + " " + sentence
                current_size = len(current_chunk.split())
        
        # Add the last chunk
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # Ensure we don't have empty chunks
        chunks = [chunk for chunk in chunks if chunk]
        
        return chunks
    
    def process_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process uploaded document: extract text, chunk, embed, and store"""
        
        try:
            # Generate unique file ID
            file_id = hashlib.md5(f"{filename}{uuid.uuid4()}".encode()).hexdigest()[:12]
            print(f"📄 Processing document: {filename} -> ID: {file_id}")
            
            # Save file
            file_extension = os.path.splitext(filename)[1].lower()
            file_path = os.path.join(self.upload_dir, f"{file_id}{file_extension}")
            
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            # Extract text
            raw_text = self.extract_text(file_path, file_extension)
            if not raw_text:
                raise ValueError("No text could be extracted from the document")
            
            cleaned_text = self.clean_text(raw_text)
            if not cleaned_text:
                raise ValueError("Text became empty after cleaning")
            
            print(f"📝 Extracted {len(cleaned_text)} characters of text")
            
            # Chunk text
            chunks = self.chunk_text(cleaned_text)
            if not chunks:
                raise ValueError("No chunks created from text")
            
            print(f"📦 Created {len(chunks)} chunks")
            
            # Generate embeddings
            print("🔄 Generating embeddings...")
            embeddings = self.embedding_model.encode(chunks)
            print(f"✅ Generated {len(embeddings)} embeddings of dimension {embeddings.shape[1]}")
            
            # Create FAISS index
            dimension = embeddings.shape[1]
            index = faiss.IndexFlatL2(dimension)
            index.add(embeddings.astype('float32'))
            print(f"✅ Created FAISS index with {index.ntotal} vectors")
            
            # Save index and chunks to disk
            index_path = os.path.join(self.index_dir, f"{file_id}.index")
            faiss.write_index(index, index_path)
            
            chunks_path = os.path.join(self.index_dir, f"{file_id}_chunks.pkl")
            with open(chunks_path, 'wb') as f:
                pickle.dump(chunks, f)
            
            # Store metadata WITH THE INDEX IN MEMORY
            self.documents[file_id] = {
                'file_id': file_id,
                'filename': filename,
                'file_path': file_path,
                'chunks': chunks,
                'num_chunks': len(chunks),
                'index': index,  # 🔴 THIS WAS MISSING - Store the actual index object
                'index_path': index_path,
                'chunks_path': chunks_path
            }
            
            print(f"✅ Document {file_id} processed and stored in memory")
            
            return {
                'file_id': file_id,
                'filename': filename,
                'chunks': len(chunks),
                'message': 'Document processed successfully'
            }
            
        except Exception as e:
            print(f"❌ Error processing document: {e}")
            import traceback
            traceback.print_exc()
            raise
    
    def load_document(self, file_id: str) -> Optional[Dict[str, Any]]:
        """Load document metadata and index"""
        
        # First check if it's already in memory
        if file_id in self.documents:
            doc = self.documents[file_id]
            # Verify the index is present
            if 'index' not in doc:
                print(f"⚠️ Document {file_id} found in memory but missing index, reloading from disk...")
                # Try to load from disk
                index_path = os.path.join(self.index_dir, f"{file_id}.index")
                if os.path.exists(index_path):
                    doc['index'] = faiss.read_index(index_path)
            return doc
        
        # If not in memory, try to load from disk
        print(f"🔄 Loading document {file_id} from disk...")
        index_path = os.path.join(self.index_dir, f"{file_id}.index")
        chunks_path = os.path.join(self.index_dir, f"{file_id}_chunks.pkl")
        
        if os.path.exists(index_path) and os.path.exists(chunks_path):
            try:
                # Load index from disk
                index = faiss.read_index(index_path)
                print(f"✅ Loaded FAISS index with {index.ntotal} vectors")
                
                # Load chunks from disk
                with open(chunks_path, 'rb') as f:
                    chunks = pickle.load(f)
                print(f"✅ Loaded {len(chunks)} chunks")
                
                # Store in memory for future use
                self.documents[file_id] = {
                    'file_id': file_id,
                    'index': index,
                    'chunks': chunks,
                    'num_chunks': len(chunks),
                    'index_path': index_path,
                    'chunks_path': chunks_path
                }
                
                print(f"✅ Document {file_id} loaded from disk and cached in memory")
                return self.documents[file_id]
                
            except Exception as e:
                print(f"❌ Error loading document {file_id} from disk: {e}")
                import traceback
                traceback.print_exc()
                return None
        
        print(f"❌ Document {file_id} not found (checked memory and disk)")
        return None
    
    def get_document_info(self, file_id: str) -> Optional[Dict]:
        """Get document information without loading the full index"""
        doc = self.load_document(file_id)
        if doc:
            return {
                'file_id': doc['file_id'],
                'chunks': len(doc['chunks']),
                'has_index': 'index' in doc,
                'in_memory': file_id in self.documents
            }
        return None
    
    def list_documents(self) -> List[Dict]:
        """List all documents in memory and on disk"""
        docs = []
        
        # Add in-memory documents
        for file_id, doc in self.documents.items():
            docs.append({
                'file_id': file_id,
                'filename': doc.get('filename', 'Unknown'),
                'chunks': len(doc['chunks']),
                'location': 'memory'
            })
        
        # Check disk for documents not in memory
        if os.path.exists(self.index_dir):
            for filename in os.listdir(self.index_dir):
                if filename.endswith('.index'):
                    file_id = filename.replace('.index', '')
                    if file_id not in self.documents:
                        chunks_path = os.path.join(self.index_dir, f"{file_id}_chunks.pkl")
                        if os.path.exists(chunks_path):
                            docs.append({
                                'file_id': file_id,
                                'filename': 'Unknown (on disk)',
                                'chunks': 'Unknown',
                                'location': 'disk'
                            })
        
        return docs


# import os
# import uuid
# import hashlib
# from typing import List, Dict, Any
# import PyPDF2
# from docx import Document
# import numpy as np
# from sentence_transformers import SentenceTransformer
# import faiss
# import pickle
# import json

# class DocumentProcessor:
#     def __init__(self, upload_dir="uploads", index_dir="vector_indices"):
#         self.upload_dir = upload_dir
#         self.index_dir = index_dir
#         os.makedirs(upload_dir, exist_ok=True)
#         os.makedirs(index_dir, exist_ok=True)
        
#         # Initialize embedding model (lightweight, runs locally)
#         self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        
#         # Store document metadata
#         self.documents = {}
        
#     def extract_text_from_pdf(self, file_path: str) -> str:
#         """Extract text from PDF file"""
#         text = ""
#         try:
#             with open(file_path, 'rb') as file:
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 for page in pdf_reader.pages:
#                     text += page.extract_text() + "\n"
#         except Exception as e:
#             print(f"Error extracting PDF: {e}")
#         return text
    
#     def extract_text_from_docx(self, file_path: str) -> str:
#         """Extract text from DOCX file"""
#         text = ""
#         try:
#             doc = Document(file_path)
#             for paragraph in doc.paragraphs:
#                 text += paragraph.text + "\n"
#         except Exception as e:
#             print(f"Error extracting DOCX: {e}")
#         return text
    
#     def extract_text_from_txt(self, file_path: str) -> str:
#         """Extract text from TXT file"""
#         try:
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 return file.read()
#         except Exception as e:
#             print(f"Error extracting TXT: {e}")
#             return ""
    
#     def extract_text(self, file_path: str, file_extension: str) -> str:
#         """Extract text based on file extension"""
#         if file_extension == '.pdf':
#             return self.extract_text_from_pdf(file_path)
#         elif file_extension == '.docx':
#             return self.extract_text_from_docx(file_path)
#         elif file_extension == '.txt':
#             return self.extract_text_from_txt(file_path)
#         else:
#             raise ValueError(f"Unsupported file type: {file_extension}")
    
#     def clean_text(self, text: str) -> str:
#         """Clean and normalize text"""
#         # Remove excessive whitespace
#         lines = [line.strip() for line in text.split('\n') if line.strip()]
#         return '\n'.join(lines)
    
#     def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 100) -> List[str]:
#         """
#         Intelligently chunk text with overlap
#         Tries to break at sentence boundaries when possible
#         """
#         chunks = []
        
#         # Split into sentences (simple approach)
#         sentences = text.replace('\n', ' ').split('. ')
        
#         current_chunk = ""
#         current_size = 0
        
#         for sentence in sentences:
#             sentence = sentence.strip() + ". "
#             sentence_size = len(sentence.split())
            
#             if current_size + sentence_size <= chunk_size:
#                 current_chunk += sentence
#                 current_size += sentence_size
#             else:
#                 if current_chunk:
#                     chunks.append(current_chunk.strip())
                
#                 # Start new chunk with overlap from previous chunk
#                 words = current_chunk.split()
#                 overlap_words = words[-overlap:] if len(words) > overlap else words
#                 current_chunk = ' '.join(overlap_words) + " " + sentence
#                 current_size = len(current_chunk.split())
        
#         # Add the last chunk
#         if current_chunk:
#             chunks.append(current_chunk.strip())
        
#         # Ensure we don't have empty chunks
#         chunks = [chunk for chunk in chunks if chunk]
        
#         return chunks
    
#     def process_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
#         """Process uploaded document: extract text, chunk, embed, and store"""
        
#         # Generate unique file ID
#         file_id = hashlib.md5(f"{filename}{uuid.uuid4()}".encode()).hexdigest()[:12]
        
#         # Save file
#         file_extension = os.path.splitext(filename)[1].lower()
#         file_path = os.path.join(self.upload_dir, f"{file_id}{file_extension}")
        
#         with open(file_path, 'wb') as f:
#             f.write(file_content)
        
#         # Extract text
#         raw_text = self.extract_text(file_path, file_extension)
#         cleaned_text = self.clean_text(raw_text)
        
#         if not cleaned_text:
#             raise ValueError("No text could be extracted from the document")
        
#         # Chunk text
#         chunks = self.chunk_text(cleaned_text)
        
#         # Generate embeddings
#         embeddings = self.embedding_model.encode(chunks)
        
#         # Create FAISS index
#         dimension = embeddings.shape[1]
#         index = faiss.IndexFlatL2(dimension)
#         index.add(embeddings.astype('float32'))
        
#         # Save index and chunks
#         index_path = os.path.join(self.index_dir, f"{file_id}.index")
#         faiss.write_index(index, index_path)
        
#         chunks_path = os.path.join(self.index_dir, f"{file_id}_chunks.pkl")
#         with open(chunks_path, 'wb') as f:
#             pickle.dump(chunks, f)
        
#         # Store metadata
#         self.documents[file_id] = {
#             'file_id': file_id,
#             'filename': filename,
#             'file_path': file_path,
#             'chunks': chunks,
#             'num_chunks': len(chunks),
#             'index_path': index_path,
#             'chunks_path': chunks_path
#         }
        
#         return {
#             'file_id': file_id,
#             'filename': filename,
#             'chunks': len(chunks),
#             'message': 'Document processed successfully'
#         }
    
#     def load_document(self, file_id: str) -> Dict[str, Any]:
#         """Load document metadata and index"""
#         if file_id not in self.documents:
#             # Try to load from disk
#             index_path = os.path.join(self.index_dir, f"{file_id}.index")
#             chunks_path = os.path.join(self.index_dir, f"{file_id}_chunks.pkl")
            
#             if os.path.exists(index_path) and os.path.exists(chunks_path):
#                 index = faiss.read_index(index_path)
#                 with open(chunks_path, 'rb') as f:
#                     chunks = pickle.load(f)
                
#                 self.documents[file_id] = {
#                     'file_id': file_id,
#                     'index': index,
#                     'chunks': chunks
#                 }
#             else:
#                 return None
        
#         return self.documents[file_id]